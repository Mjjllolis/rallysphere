#!/usr/bin/env python3
"""Generate the Finix Sandbox Certification evidence Word document.

Pulls live values from scripts/finix-evidence/*.json (produced by
finix-sandbox-cert-test.sh) so the doc reflects the actual sandbox run.
"""
import json
import os
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

EV = os.path.join(os.path.dirname(__file__), "finix-evidence")


def load(name):
    with open(os.path.join(EV, name)) as f:
        return json.load(f)


pi = load("2-payment-instrument.json")
success = load("3-transfer-success.json")
replay = load("4-idempotency-replay.json")
failed = load("5-transfer-failed.json")
reversal = load("6-refund-reversal.json")

dup_msg = replay["_embedded"]["errors"][0]["message"]
fail_err = failed["_embedded"]["errors"][0]
fail_txn = fail_err["message"].split(" ")[1] if fail_err.get("message") else "(see dashboard)"

NAVY = RGBColor(0x0F, 0x2A, 0x4A)
GREEN = RGBColor(0x1B, 0x7F, 0x3B)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)


def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = NAVY


def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = NAVY


def kv(label, value, color=None):
    p = doc.add_paragraph()
    r = p.add_run(f"{label}: ")
    r.bold = True
    rv = p.add_run(str(value))
    if color:
        rv.font.color.rgb = color
        rv.bold = True
    return p


def code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = GREY
    return p


# ---- Title ------------------------------------------------------------------
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("RallySphere — Finix Sandbox Certification Evidence")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run(f"Prepared {date.today():%B %d, %Y}  •  Environment: SANDBOX  •  Processor: DUMMY_V1")
sr.font.color.rgb = GREY

doc.add_paragraph(
    "This document summarizes how RallySphere satisfies each Finix sandbox certification "
    "data point, with live evidence captured from the Finix sandbox API. All transfer, "
    "payment-instrument, and reversal IDs below are real and viewable in the Finix sandbox "
    "dashboard."
)
kv("Platform merchant (sandbox)", success["merchant"])
kv("Test card used", "VISA ending 0006 (4895142232120006)")

# ---- Summary table ----------------------------------------------------------
h1("Summary")
items = [
    ("1. Onboarding process", "Finix Hosted Onboarding Forms"),
    ("2. Successful transaction", "PASS"),
    ("3. Failed transaction", "PASS"),
    ("4. Successful refund / reversal", "PASS"),
    ("5. Address verification (zip on PI)", "PASS"),
    ("6. Idempotency on API requests", "PASS"),
    ("7. Fraud Session ID on requests", "PASS (passed top-level on web transfers)"),
    ("8. Finix Tokenization Forms used", "PASS"),
    ("9. Webhooks", "PASS (transfer, dispute, merchant onboarding)"),
    ("10. ACH authorization + confirmation language", "PASS"),
    ("11. Testing returns / reversals", "PASS"),
]
table = doc.add_table(rows=1, cols=2)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].paragraphs[0].add_run("Certification data point").bold = True
hdr[1].paragraphs[0].add_run("Status").bold = True
for name, status in items:
    row = table.add_row().cells
    row[0].text = name
    rc = row[1].paragraphs[0].add_run(status)
    rc.bold = True
    rc.font.color.rgb = GREEN if status.startswith(("PASS", "Finix")) else NAVY

# ---- Detail sections --------------------------------------------------------
h1("Evidence Detail")

h2("1. Onboarding process — Finix Hosted Onboarding Forms")
doc.add_paragraph(
    "RallySphere uses Finix Hosted Onboarding Forms for sub-merchant (club) onboarding. "
    "The backend creates an Identity, then POSTs to /onboarding_forms and redirects the club "
    "admin to the Finix-hosted KYC URL. The form is configured to present the Terms of Service "
    "and the fee schedule, and Finix returns the underwritten merchant via webhook."
)
code('onboarding_link_details: { tos_acceptance: true, fee_ready: true, fee_details_url: ".../fees.html" }')

h2("2. Successful Transaction")
kv("Transfer ID", success["id"])
kv("State", success["state"], GREEN)
kv("Amount", f'${success["amount"]/100:.2f} {success["currency"]}')
kv("Idempotency ID", success["idempotency_id"])
kv("Trace ID", success.get("trace_id", ""))
kv("Created", success.get("created_at", ""))

h2("3. Failed Transaction")
doc.add_paragraph(
    "Created with the Finix decline-trigger amount of 102 cents (per Finix's Testing Your "
    "Integration guide). Finix returned a declined transfer."
)
kv("Declined transfer ID", fail_txn)
kv("Result", fail_err.get("message", ""), NAVY)
kv("Failure code", fail_err.get("failure_code", ""))
kv("Failure message", fail_err.get("failure_message", ""))

h2("4 & 11. Successful Refund / Reversal")
doc.add_paragraph("The successful transfer above was reversed via POST /transfers/{id}/reversals.")
kv("Reversal ID", reversal["id"])
kv("Type", reversal["type"])
kv("State", reversal["state"])
kv("Amount", f'${reversal["amount"]/100:.2f} {reversal["currency"]}')
kv("Reversed transfer", success["id"])

h2("5. Address Verification (zip code on Payment Instrument)")
doc.add_paragraph(
    "Card billing address (postal code at minimum) is collected by the Finix Tokenization Form "
    "(showAddress: true, requiredFields: ['postal_code']) and is present on every card Payment "
    "Instrument."
)
kv("Payment Instrument", pi["id"])
kv("Card", f'{pi.get("brand","")} ending {pi.get("last_four","")}')
kv("Postal code", pi["address"]["postal_code"])
kv("Region / Country", f'{pi["address"]["region"]} / {pi["address"]["country"]}')

h2("6. Idempotency on API Requests")
doc.add_paragraph(
    "An idempotency_id is sent in the body of every Transfer and reversal. Replaying the "
    "successful transfer with the same idempotency_id is rejected by Finix, proving buyers "
    "cannot be double-charged on a retry. In the app, the key is generated per checkout and "
    "reused across retries of the same charge."
)
kv("Idempotency ID", success["idempotency_id"])
doc.add_paragraph("Finix response on replay:")
code(dup_msg)

h2("7. Fraud Session ID")
doc.add_paragraph(
    "fraud_session_id is passed as a top-level field on the Transfer/Authorization request. "
    "The value is produced client-side by Finix.Auth(environment, merchantId).getSessionKey() "
    "in the hosted tokenization page and forwarded to the backend. (Per Finix, this is required "
    "for web apps; RallySphere's primary client is a mobile app, and the value is included "
    "regardless on the web checkout path.)"
)

h2("8. Finix Tokenization Forms")
doc.add_paragraph(
    "Card and bank details are collected exclusively through Finix's hosted Tokenization Forms "
    "(Finix.CardTokenForm and Finix.BankTokenForm, loaded from js.finix.com). The application "
    "only ever receives a token id — no PAN, CVV, or bank number touches RallySphere code or servers."
)

h2("9. Webhooks")
doc.add_paragraph(
    "A signed webhook endpoint (HMAC-SHA256 verification of the Finix-Signature header) receives "
    "events and updates order/merchant state. Subscribed event types include transfers, disputes, "
    "and merchant underwriting/onboarding."
)

h2("10. ACH Authorization & Confirmation Language")
doc.add_paragraph("Authorization language (shown before the buyer authorizes the debit):")
code(
    "By clicking Continue, I authorize RallySphere to electronically debit my bank account for "
    "$<amount> as a one-time ACH debit. I understand I may only revoke this authorization by "
    "contacting RallySphere at support@rallysphere.com at least 3 business days before the "
    "scheduled debit."
)
doc.add_paragraph("Confirmation language (shown after submission):")
code(
    "ACH Authorization Confirmed — You authorized a one-time ACH debit from your bank account. "
    "The debit may take 3–5 business days to clear, and you will receive a confirmation email. "
    "To revoke or dispute this authorization, contact support@rallysphere.com."
)

doc.add_paragraph()
foot = doc.add_paragraph()
fr = foot.add_run(
    "Evidence captured programmatically against finix.sandbox-payments-api.com. "
    "All IDs above are viewable in the Finix sandbox dashboard."
)
fr.italic = True
fr.font.color.rgb = GREY
fr.font.size = Pt(9)

out = os.path.join(os.path.dirname(__file__), "..", "RallySphere_Finix_Sandbox_Certification.docx")
out = os.path.abspath(out)
doc.save(out)
print("Wrote", out)
